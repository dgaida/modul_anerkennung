"""
Erzeugt ein Word-Dokument mit einer Äquivalenzliste zwischen zwei Prüfungsordnungen.

Dieses Skript liest die Äquivalenzliste aus `data/aequivalenzliste.md` und ergänzt sie
um alle weiteren Pflichtmodule aus den angegebenen Prüfungsordnungen (Standard: inf_inf2 und inf_inf3).
Das Ergebnis ist eine 4-spaltige Tabelle (Alte PO Name, Alte PO ECTS, Neue PO Name, Neue PO ECTS),
die primär nach dem Semester der neuen PO sortiert ist. Äquivalente Module stehen in derselben Zeile.

Nutzung:
    PYTHONPATH=. python3 scripts/create_equivalence_table_word.py --old-po inf_inf2 --new-po inf_inf3
"""

import argparse
import json
import logging
import os
import sys
import urllib.error
import urllib.request
from typing import Any, Dict, List, Set

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from dotenv import load_dotenv

# Environment laden
load_dotenv()

# Logging konfigurieren
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger("create_word_table")

# Konfiguration und Environment laden
try:
    import modul_anerkennung.config  # noqa: E402, F401
except ImportError:
    logger.warning("Konnte modul_anerkennung.config nicht laden. Stelle sicher, dass PYTHONPATH gesetzt ist.")

BASE_URL = "https://module.gm.th-koeln.de/api"
API_TOKEN = os.environ.get("MOCOGI_API_TOKEN") or os.environ.get("MOCOGI_API_KEY")

def api_call(url: str, method: str = "GET", data: Dict[str, Any] = None) -> Any:
    """Führt einen API-Aufruf aus und handhabt Authentifizierungsfehler.

    Args:
        url (str): Die URL des API-Endpunkts.
        method (str): Die HTTP-Methode (Standard: "GET").
        data (Dict[str, Any], optional): Optionales Dictionary mit JSON-Daten für den Body.

    Returns:
        Any: Die Antwort der API als Dictionary oder Liste.
    """
    headers = {
        "Content-Type": "application/json",
    }
    if API_TOKEN:
        headers["Authorization"] = f"Bearer {API_TOKEN}"

    req = urllib.request.Request(url, headers=headers, method=method)
    if data:
        req.data = json.dumps(data).encode("utf-8")

    try:
        with urllib.request.urlopen(req) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        if e.code in [401, 403]:
            logger.error("Authentifizierungsfehler (401/403). Prüfen Sie Ihren MOCOGI_API_TOKEN.")
        else:
            logger.error(f"HTTP-Fehler beim API-Aufruf ({url}): {e.code} {e.reason}")
        return None
    except Exception as e:
        logger.error(f"Unerwarteter Fehler beim API-Aufruf ({url}): {e}")
        raise

def get_module_metadata(m: Dict[str, Any]) -> Dict[str, Any]:
    """Extrahiert Metadaten robust aus verschiedenen Verschachtelungsebenen.

    Args:
        m (Dict[str, Any]): Das Modul- oder Draft-Objekt von der API.

    Returns:
        Dict[str, Any]: Das Dictionary mit den Metadaten.
    """
    # Priorität: top-level metadata > module.metadata > module (als metadata)
    metadata = m.get("metadata") or m.get("module", {}).get("metadata") or m.get("module") or {}
    if not isinstance(metadata, dict):
        return {}
    return metadata

def get_module_title(m: Dict[str, Any]) -> str:
    """Extrahiert den Titel aus einem Modul- oder Draft-Objekt.

    Args:
        m (Dict[str, Any]): Das Modul- oder Draft-Objekt.

    Returns:
        str: Der Titel des Moduls.
    """
    meta = get_module_metadata(m)
    title = meta.get("title") or m.get("title") or ""
    return str(title).strip()

def get_module_ects(m: Dict[str, Any]) -> str:
    """Extrahiert die ECTS aus einem Modul- oder Draft-Objekt.

    Args:
        m (Dict[str, Any]): Das Modul- oder Draft-Objekt.

    Returns:
        str: Die ECTS-Punkte als String.
    """
    meta = get_module_metadata(m)
    # ECTS stehen bei Drafts oft auf Top-Level, bei Modulen in Metadata
    ects = m.get("ects") or meta.get("ects") or ""
    return str(ects)

def is_mandatory(m: Dict[str, Any], po_id: str) -> bool:
    """Prüft, ob ein Modul für eine bestimmte PO ein Pflichtmodul ist.

    Args:
        m (Dict[str, Any]): Das Modul- oder Draft-Objekt.
        po_id (str): Die Kennung der Prüfungsordnung (z.B. 'inf_inf2').

    Returns:
        bool: True, wenn es ein Pflichtmodul ist, sonst False.
    """
    # Fall 1: Draft-Objekt
    mandatory_pos = m.get("mandatoryPOs") or []
    if po_id in mandatory_pos:
        return True

    # Fall 2: Publiziertes Modul - Metadaten Struktur
    meta = get_module_metadata(m)
    po_info = meta.get("po") or {}
    mandatory_list = po_info.get("mandatory") or []
    for item in mandatory_list:
        if item.get("po") == po_id:
            return True

    return False

def get_semester(m: Dict[str, Any], po_id: str) -> int:
    """Extrahiert das empfohlene Semester für eine bestimmte PO.

    Args:
        m (Dict[str, Any]): Das Modul- oder Draft-Objekt.
        po_id (str): Die Kennung der Prüfungsordnung.

    Returns:
        int: Das Semester als Integer oder 99 als Fallback.
    """
    meta = get_module_metadata(m)
    po_info = meta.get("po") or {}

    # Prüfe mandatory und optional Listen
    for category in ["mandatory", "optional"]:
        items = po_info.get(category) or []
        for item in items:
            if item.get("po") == po_id:
                sem_list = item.get("recommendedSemester") or []
                if sem_list:
                    try:
                        return int(sem_list[0])
                    except (ValueError, IndexError):
                        continue

    # Fallback für Drafts (manchmal direkt im Objekt)
    # In Drafts ist die Struktur oft anders, falls oben nicht gefunden:
    return 99

def parse_markdown_table(file_path: str) -> List[Dict[str, Any]]:
    """Parst die Äquivalenzliste aus einer Markdown-Datei.

    Args:
        file_path (str): Pfad zur MD-Datei.

    Returns:
        List[Dict[str, Any]]: Liste von Dictionaries mit 'source', 'target' und 'semester'.
    """
    entries = []
    if not os.path.exists(file_path):
        logger.error(f"Datei nicht gefunden: {file_path}")
        return entries

    with open(file_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line or not line.startswith("|") or "Modul in PO" in line or "| :---" in line:
            continue

        parts = [p.strip() for p in line.split("|") if p.strip()]
        if len(parts) >= 2:
            source_title = parts[0]
            target_title = parts[1]
            semester = 0
            if len(parts) >= 3:
                try:
                    semester = int(parts[2])
                except ValueError:
                    semester = 0

            entries.append({
                "source": source_title,
                "target": target_title,
                "semester": semester
            })
    return entries

def main() -> None:
    """Hauptfunktion zur Tabellenerzeugung."""
    parser = argparse.ArgumentParser(description="Erzeugt ein Word-Dokument mit einer Äquivalenzliste.")
    parser.add_argument("--old-po", default="inf_inf2", help="ID der alten Prüfungsordnung (Default: inf_inf2)")
    parser.add_argument("--new-po", default="inf_inf3", help="ID der neuen Prüfungsordnung (Default: inf_inf3)")
    args = parser.parse_args()

    old_po = args.old_po
    new_po = args.new_po

    logger.info(f"Starte Generierung der Äquivalenzliste für Word: {old_po} -> {new_po}")

    # 1. Daten von API abrufen (alle aktiven für Lookup)
    logger.info(f"Lade Module der alten PO ({old_po})...")
    old_po_all = api_call(f"{BASE_URL}/modules?po={old_po}&active=true&select=metadata") or []
    
    logger.info(f"Lade Entwürfe/Module der neuen PO ({new_po})...")
    # Zuerst Entwürfe prüfen (üblich für neue POs)
    drafts_resp = api_call(f"{BASE_URL}/moduleDrafts")
    new_po_all = []
    if drafts_resp:
        all_drafts = (drafts_resp.get("direct") or []) + (drafts_resp.get("indirect") or [])
        new_po_all = all_drafts

    # Auch publizierte Module der neuen PO laden (falls vorhanden)
    published_new = api_call(f"{BASE_URL}/modules?po={new_po}&active=true&select=metadata") or []
    new_po_all.extend(published_new)

    if not old_po_all and not new_po_all:
        logger.warning("Keine Module für beide Prüfungsordnungen gefunden. Prüfen Sie die IDs.")

    # 2. Äquivalenzliste einlesen
    list_path = "data/aequivalenzliste.md"
    equiv_list = parse_markdown_table(list_path)
    logger.info(f"{len(equiv_list)} Einträge aus {list_path} geladen.")

    # 3. Hilfs-Mappings erstellen (Titel -> Modul-Objekt)
    # Wir filtern hier auf Pflichtmodule
    old_by_title = {get_module_title(m).lower(): m for m in old_po_all if is_mandatory(m, old_po)}
    new_by_title = {get_module_title(m).lower(): m for m in new_po_all if is_mandatory(m, new_po)}

    used_old: Set[str] = set()
    used_new: Set[str] = set()
    rows: List[Dict[str, Any]] = []

    # 4. Einträge aus der Äquivalenzliste verarbeiten (Zusammenführung)
    for eq in equiv_list:
        src_t_raw = eq["source"]
        tgt_t_raw = eq["target"]
        src_t = src_t_raw.lower()
        tgt_t = tgt_t_raw.lower()

        m_old = old_by_title.get(src_t)
        m_new = new_by_title.get(tgt_t)

        # Filter: Mindestens eines muss mandatory sein
        is_old_man = m_old and is_mandatory(m_old, old_po)
        is_new_man = m_new and is_mandatory(m_new, new_po)

        if not is_old_man and not is_new_man:
            logger.debug(f"Überspringe Wahl-Äquivalenz: {src_t_raw} -> {tgt_t_raw}")
            continue

        if is_old_man:
            used_old.add(src_t)
        if is_new_man:
            used_new.add(tgt_t)

        # Bestimme Sortier-Semester (Bevorzugt neue PO)
        sort_sem = eq["semester"] or (get_semester(m_new, new_po) if m_new else get_semester(m_old, old_po))

        rows.append({
            "old": m_old,
            "old_title_fallback": src_t_raw if not m_old else None,
            "new": m_new,
            "new_title_fallback": tgt_t_raw if not m_new else None,
            "sort_semester": sort_sem,
            "sort_title": (get_module_title(m_new) if m_new else (get_module_title(m_old) if m_old else tgt_t_raw))
        })

    # 5. Restliche Pflichtmodule hinzufügen (ohne Äquivalent in der Liste)
    for title_lower, m in old_by_title.items():
        if title_lower not in used_old:
            rows.append({
                "old": m,
                "new": None,
                "sort_semester": get_semester(m, old_po),
                "sort_title": get_module_title(m)
            })

    # 6. Restliche Pflichtmodule hinzufügen (ohne Äquivalent in der Liste)
    for title_lower, m in new_by_title.items():
        if title_lower not in used_new:
            rows.append({
                "old": None,
                "new": m,
                "sort_semester": get_semester(m, new_po),
                "sort_title": get_module_title(m)
            })

    # 7. Sortierung: nach Semester, dann alphabetisch nach Titel
    print(f"INFO: Die Sortierung erfolgt primär nach dem empfohlenen Semester der neuen Prüfungsordnung ({new_po}).")
    rows.sort(key=lambda x: (x["sort_semester"], x["sort_title"].lower()))

    # 8. Word-Dokument erstellen
    doc = Document()

    # Seitenränder für Tabellen optimieren
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    doc.add_heading(f'Anlage: Äquivalenzliste {old_po.upper()} / {new_po.upper()}', 0)

    # Tabelle mit 4 Spalten erzeugen
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header setzen
    hdr_cells = table.rows[0].cells
    header_labels = [f'Modul {old_po.upper()}', f'ECTS {old_po.upper()}', f'Modul {new_po.upper()}', f'ECTS {new_po.upper()}']
    for i, label in enumerate(header_labels):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Datenzeilen befüllen
    for row_data in rows:
        row_cells = table.add_row().cells

        # Alte PO Spalten (1 & 2)
        m_old = row_data["old"]
        if m_old:
            row_cells[0].text = get_module_title(m_old)
            row_cells[1].text = get_module_ects(m_old)
        elif row_data.get("old_title_fallback"):
            row_cells[0].text = row_data["old_title_fallback"]
            row_cells[1].text = "-"
        else:
            row_cells[0].text = ""
            row_cells[1].text = ""

        # Neue PO Spalten (3 & 4)
        m_new = row_data["new"]
        if m_new:
            row_cells[2].text = get_module_title(m_new)
            row_cells[3].text = get_module_ects(m_new)
        elif row_data.get("new_title_fallback"):
            row_cells[2].text = row_data["new_title_fallback"]
            row_cells[3].text = "-"
        else:
            row_cells[2].text = ""
            row_cells[3].text = ""

        # ECTS-Spalten zentrieren
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 9. Studiengang- und PO-Namen ermitteln
    study_programs = api_call(f"{BASE_URL}/studyPrograms") or []

    # Defaults/Fallbacks ermitteln
    import re
    def get_fallback_version(po_id: str) -> str:
        match = re.search(r'(\d+)', po_id)
        if match:
            return f"PO {match.group(1)}"
        return po_id.upper()

    def get_fallback_studiengang(po_id: str) -> str:
        if po_id.startswith("inf_wi"):
            return "Wirtschaftsinformatik"
        elif po_id.startswith("inf_mim"):
            return "Medieninformatik"
        elif po_id.startswith("inf_wsc"):
            return "Web Science"
        return "Informatik"

    old_studiengang = get_fallback_studiengang(old_po)
    old_po_version_str = get_fallback_version(old_po)
    new_studiengang = get_fallback_studiengang(new_po)
    new_po_version_str = get_fallback_version(new_po)

    for program in study_programs:
        prog_po = program.get("po", {})
        prog_po_id = prog_po.get("id")
        if prog_po_id == old_po:
            old_studiengang = program.get("deLabel") or old_studiengang
            version = prog_po.get("version")
            if version is not None:
                old_po_version_str = f"PO {version}"
        if prog_po_id == new_po:
            new_studiengang = program.get("deLabel") or new_studiengang
            version = prog_po.get("version")
            if version is not None:
                new_po_version_str = f"PO {version}"

    # ECTS-Summen berechnen
    def sum_ects_for_po(by_title_dict: Dict[str, Any]) -> int:
        total = 0
        for m in by_title_dict.values():
            ects_str = get_module_ects(m)
            try:
                if ects_str:
                    total += int(float(ects_str))
            except (ValueError, TypeError):
                pass
        return total

    sum1 = sum_ects_for_po(old_by_title)
    sum2 = sum_ects_for_po(new_by_title)

    # Letzte Zeile für Summen hinzufügen
    sum_row_cells = table.add_row().cells
    sum_texts = [
        f"Summe Studiengang {old_studiengang} ({old_po_version_str})",
        str(sum1),
        f"Summe Studiengang {new_studiengang} ({new_po_version_str})",
        str(sum2)
    ]

    for i, text in enumerate(sum_texts):
        p = sum_row_cells[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        if i in [1, 3]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_file = f"aequivalenzliste_{old_po}_{new_po}.docx"
    doc.save(output_file)
    logger.info(f"Dokument erfolgreich unter '{output_file}' gespeichert.")
    print(f"\nERFOLG: {output_file} mit {len(rows)} Zeilen erstellt.")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        logger.error(f"Fehler bei der Ausführung: {e}")
        sys.exit(1)
